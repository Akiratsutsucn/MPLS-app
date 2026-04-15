from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload
from io import BytesIO
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.database import get_db
from app.models.project import Project
from app.models.checklist import ChecklistItem
from app.models.document import Document
from app.models.log import ProjectLog
from app.models.user import User
from app.auth import get_current_user
from app.seed import DEFAULT_CHECKLIST

router = APIRouter(prefix="/projects", tags=["export"])

PHASE_LABELS = {
    "grading": "定级",
    "filing": "备案",
    "rectification": "整改",
    "evaluation": "测评",
    "supervision": "监督检查",
}


@router.get("/{project_id}/export")
async def export_project(
    project_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(Project)
        .options(
            selectinload(Project.assignee),
            selectinload(Project.creator),
            selectinload(Project.checklist_items),
            selectinload(Project.documents).selectinload(Document.uploader),
        )
        .where(Project.id == project_id)
    )
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    log_result = await db.execute(
        select(ProjectLog)
        .options(selectinload(ProjectLog.author))
        .where(ProjectLog.project_id == project_id)
        .order_by(ProjectLog.created_at.desc())
        .limit(20)
    )
    logs = log_result.scalars().all()

    doc = DocxDocument()

    # Title
    title = doc.add_heading(f"等级保护项目报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Project info section
    doc.add_heading("项目基本信息", level=1)
    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = "Table Grid"

    def add_info_row(label: str, value: str):
        row = info_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value

    add_info_row("项目名称", project.name)
    add_info_row("系统名称", project.system_name)
    add_info_row("等保等级", f"第{project.level}级")
    add_info_row("当前阶段", PHASE_LABELS.get(project.phase, project.phase))
    add_info_row("委托单位", project.client_org or "-")
    add_info_row("负责人", project.assignee.name if project.assignee else "-")
    add_info_row("截止日期", project.deadline.strftime("%Y-%m-%d") if project.deadline else "-")
    add_info_row("创建人", project.creator.name if project.creator else "-")
    add_info_row("创建时间", project.created_at.strftime("%Y-%m-%d %H:%M"))
    if project.notes:
        add_info_row("备注", project.notes)

    # Checklist by phase
    doc.add_heading("各阶段检查清单", level=1)
    items_by_phase: dict[str, list[ChecklistItem]] = {}
    for item in project.checklist_items:
        items_by_phase.setdefault(item.phase, []).append(item)

    for phase_key in DEFAULT_CHECKLIST.keys():
        phase_items = sorted(items_by_phase.get(phase_key, []), key=lambda x: x.sort_order)
        if not phase_items:
            continue

        doc.add_heading(PHASE_LABELS.get(phase_key, phase_key), level=2)
        cl_table = doc.add_table(rows=1, cols=3)
        cl_table.style = "Table Grid"
        header_cells = cl_table.rows[0].cells
        header_cells[0].text = "序号"
        header_cells[1].text = "检查项"
        header_cells[2].text = "完成状态"

        for idx, item in enumerate(phase_items, start=1):
            row = cl_table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = item.content
            row.cells[2].text = "已完成" if item.is_completed else "未完成"

    # Documents section
    doc.add_heading("文档列表", level=1)
    if project.documents:
        doc_table = doc.add_table(rows=1, cols=4)
        doc_table.style = "Table Grid"
        hdr = doc_table.rows[0].cells
        hdr[0].text = "文件名"
        hdr[1].text = "阶段"
        hdr[2].text = "上传人"
        hdr[3].text = "上传时间"

        for d in sorted(project.documents, key=lambda x: x.uploaded_at, reverse=True):
            row = doc_table.add_row()
            row.cells[0].text = d.filename
            row.cells[1].text = PHASE_LABELS.get(d.phase, d.phase)
            row.cells[2].text = d.uploader.name if d.uploader else "-"
            row.cells[3].text = d.uploaded_at.strftime("%Y-%m-%d %H:%M")
    else:
        doc.add_paragraph("暂无文档")

    # Recent logs section
    doc.add_heading("最近操作记录", level=1)
    if logs:
        log_table = doc.add_table(rows=1, cols=3)
        log_table.style = "Table Grid"
        lhdr = log_table.rows[0].cells
        lhdr[0].text = "时间"
        lhdr[1].text = "操作人"
        lhdr[2].text = "内容"

        for log in logs:
            row = log_table.add_row()
            row.cells[0].text = log.created_at.strftime("%Y-%m-%d %H:%M")
            row.cells[1].text = log.author.name if log.author else "系统"
            row.cells[2].text = log.content
    else:
        doc.add_paragraph("暂无记录")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = project.name.replace(" ", "_")
    filename = f"{safe_name}_等保报告.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
